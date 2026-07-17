# -*- coding: utf-8 -*-
import copy
import io
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"C:\Users\mynam\Desktop\인하대학교\KAIST AT X 실패 공모전"
DOCX_PATH = ROOT + r"\제안서_수정본_사본.docx"

log = io.open(ROOT + r"\_workspace\patch_log.txt", "w", encoding="utf-8")

doc = Document(DOCX_PATH)

p11 = doc.paragraphs[11]
run0 = p11.runs[0]
text = run0.text
marker = "삼고 있지만"
log.write("marker in text: %s\n" % (marker in text))
idx = text.index(marker) + len(marker)
before = text[:idx]
after = text[idx:]
log.write("before=%r\n" % before)
log.write("after=%r\n" % after)

run0.text = before


def clone_run_props(src_r_elem):
    new_r = OxmlElement("w:r")
    rpr = src_r_elem.find(qn("w:rPr"))
    if rpr is not None:
        new_r.append(copy.deepcopy(rpr))
    return new_r


def set_text(r_elem, text):
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r_elem.append(t)


citation_r = clone_run_props(run0._r)
set_text(
    citation_r,
    "(12~21세 응답자 중 정신건강 고민을 AI 챗봇에게 먼저 구한 비율이 8명 중 1명에 달한다는 조사도 있다",
)
run0._r.addnext(citation_r)

sup_r = clone_run_props(run0._r)
rpr = sup_r.find(qn("w:rPr"))
if rpr is None:
    rpr = OxmlElement("w:rPr")
    sup_r.insert(0, rpr)
vertalign = OxmlElement("w:vertAlign")
vertalign.set(qn("w:val"), "superscript")
rpr.append(vertalign)
set_text(sup_r, "2")
citation_r.addnext(sup_r)

closeparen_r = clone_run_props(run0._r)
set_text(closeparen_r, ")")
sup_r.addnext(closeparen_r)

rest_r = clone_run_props(run0._r)
set_text(rest_r, after)
closeparen_r.addnext(rest_r)

# Accent border under title
title_p = doc.paragraphs[0]
pPr = title_p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "18")
bottom.set(qn("w:space"), "8")
bottom.set(qn("w:color"), "E4573D")
pBdr.append(bottom)
pPr.append(pBdr)

# Second footnote
last_p = doc.paragraphs[-1]
new_p = doc.add_paragraph()
new_p.paragraph_format.space_before = last_p.paragraph_format.space_before
r = new_p.add_run(
    '² Brown University School of Public Health (2025.11.18), "One in eight adolescents and young adults use AI chatbots for mental health advice."'
)
r.font.size = last_p.runs[0].font.size
r.font.bold = False
r.font.italic = False
r.font.name = last_p.runs[0].font.name
r.element.rPr.rFonts.set(qn("w:eastAsia"), last_p.runs[0].font.name)

doc.save(DOCX_PATH)
log.write("Saved.\n")
log.write("New paragraph 11 text: %s\n" % p11.text)
log.write("Total paragraphs now: %d\n" % len(doc.paragraphs))
log.close()
print("done")
