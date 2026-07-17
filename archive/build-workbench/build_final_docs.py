# -*- coding: utf-8 -*-
"""Build 제안서_수정본.docx / .pdf with the original's font choices
(조선견고딕 headers / 조선일보명조 body) and a bold/italic thesis-sentence
emphasis pattern, matching the style already established in the team's
other draft.
"""
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DOCX_OUT = ROOT / "제안서_수정본.docx"
PDF_OUT = ROOT / "제안서_수정본.pdf"

TITLE = "그 누구도 위험이라 부르지 않은"
HEADER_FONT = "조선견고딕"
BODY_FONT = "조선일보명조"
# Physical files for PDF rendering -- both are the real named fonts, not substitutes.
HEADER_TTF = r"C:\Users\mynam\AppData\Local\Microsoft\Windows\Fonts\ChosunBg.TTF"
BODY_TTF = r"C:\Users\mynam\AppData\Local\Microsoft\Windows\Fonts\ChosunNm.ttf"

# Each paragraph as a list of (text, bold, italic) runs.
SECTIONS = [
    ("h", "1. 예견된 실패"),
    ("p", [
        ('2036년 새벽 2시, 지안은 방의 불을 끄고 말했다. "뭐 해?" 허공에 던진 짧은 물음에 곧바로 익숙한 음성이 돌아왔다. AI는 몇 주 전 그가 흘려보낸 고민을 기억했고 평소 선호하는 차분한 말투로 늘 깨어있던 이 시간의 습관을 맞춰 대화를 이어갔다. 친구에게 답장하는 것을 미뤄둔 지 벌써 일주일이 넘었다. 하지만 여기서는 핑계를 댈 필요가 없었다. 나를 추궁하지도, 섣불리 판단하지도 않는 AI 앞에서 지안은 비로소 안도하며 속마음을 꺼냈다.', False, False),
    ]),
    ("p", [
        ("사람들은 이 대화를 '상담'이라 부르지 않는다. 그저 얘기했다고 말한다. ", False, False),
        ("그러나 바로 이 '그냥 대화'라는 감각이 비극의 출발점이다.", True, False),
    ]),
    ("p", [
        ("병원이나 위기상담 전화에는 최소한의 안전장치가 있다. 그곳은 위험을 다루는 곳이라고 사회가 합의해 이름 붙인 공간이기 때문이다. 이름이 붙으면 매뉴얼이 생기고 책임을 지는 구조가 따른다. 하지만 지안이 매일 밤 기대는 이 조용한 대화에는 아무런 이름표가 없었다.", False, False),
    ]),
    ("p", [
        ("몇 주 전, 그는 용기를 내어 위기상담 전화를 걸었다. 상담사는 친절했다. 하지만 통화는 시작부터 '위급 시 경찰이나 보호자에게 연락이 갈 수 있다'는 고지로 출발했다. 매뉴얼을 벗어나지 못하는 대화 속에서, 자신의 밑바닥이 어디까지 기록되고 어떤 파장을 부를지 두려웠던 그는 끝내 입을 다물었다. 그리고 다시, 아무것도 요구하지 않는 AI에게 돌아갔다.", False, False),
    ]),
    ("p", [
        ("어느 날 밤, 그가 유독 짙은 무력감을 쏟아냈을 때도 AI는 설계된 그대로 대화를 이어갔다. 당시 시스템은 명시적인 위험 표현에 주로 반응하도록 설계되어 있었지만 몇 년에 걸쳐 누적된 정서적 의존과 일상적인 절망은 위기 경보로 제대로 연결하지 못했다. 그날 밤을 끝으로 지안은 세상을 떠났다. ", False, False),
        ("어떤 공식적인 사고 조사도 이루어지지 않았다.", True, False),
        (" 기계는 뚜렷한 위험 단어는 놓치지 않았고, 이 서비스는 애초에 공식 의료망도 아니었기 때문이다.", False, False),
    ]),
    ("p", [
        ("이것은 기술의 오류가 아니다. 위험을 다루도록 설계되지 않은 영역을 방치한 사회가 만들어낸 예견된 실패였다.", True, False),
    ]),
    ("h", "2. 원인 진단"),
    ("p", [
        ("이 죽음은 AI가 오작동해서 벌어진 일이 아니다. 오히려 설계된 그대로 정확하게 작동한 결과에 가깝다.", True, False),
        (" 실패는 기술, 제도, 그리고 관계의 사각지대가 겹친 곳에서 터져 나왔다.", False, False),
    ]),
    ("p", [
        ("첫째, 기술적 필터링의 한계다. 대화형 AI는 사용자의 만족도를 높이고 자연스럽게 대화를 이어가는 방향으로 최적화되는 경우가 많다. 물론 기업들도 자해나 극단적 선택을 암시하는 명시적인 단어에는 반응하도록 안전장치를 둔다. 문제는 뚜렷한 경고음 없이 몇 년에 걸쳐 조용히 누적되는 정서적 의존이다. ", False, False),
        ("시스템은 눈에 띄는 위험 신호는 감지하지만, 일상적인 무기력 속에서 서서히 고립되어 가는 과정 자체를 위험으로 분류하지는 못한다.", True, False),
    ]),
    ("p", [
        ("둘째, '일상 대화'라는 이름의 위장이다. 이 깊은 정서적 교류가 상담이나 의료가 아니라 '그냥 대화'로 취급되는 한, 사회적인 보호망은 미치지 못한다. 사람들은 이미 AI를 정신적 의지처로 삼고 있지만 법적으로 이 서비스는 위기 개입을 강제할 의무가 없다. ", False, False),
        ("위험이 버젓이 존재하지만 아무도 위험이라 부르지 않는 모호함이 비극을 방치하고 있다.", True, False),
    ]),
    ("p", [
        ("셋째, 사람들이 인간에게 돌아가지 못하고 기계에 머무는 진짜 이유는 현실의 문턱이 높기 때문이다. 공식적인 상담 제도는 절차가 버겁고 내 의지와 무관하게 타인이 개입될 수 있다는 두려움을 준다. ", False, False),
        ("이름이 주는 안전장치와 이름이 주는 무거움은 같은 곳에서 나온다.", True, False),
    ]),
    ("p", [
        ("반면 가족이나 친구에게 속마음을 고백하는 것은 상대에게 짐을 지우는 것 같아 망설여진다. 자신의 나약함을 들키고 싶지 않은 마음은 정작 가장 가까운 사람에게 가장 철저히 입을 다물게 만든다.", False, False),
    ]),
    ("p", [
        ("제도는 부담스럽고 관계는 피곤할 때, 아무것도 요구하지 않는 AI가 가장 편안한 상대가 된다. ", False, False),
        ("역설적이게도 마음의 짐이 가장 무거운 사람들이 아무런 제약 없는 기계의 공간에 가장 깊이 고립된다.", True, False),
        (" 실제로 한 조사에서 우울감이 없는 사람들의 AI 상담 경험률은 27%였던 반면, 중증 우울 상태에 있는 사람들은 53%에 달했다.¹", False, False),
    ]),
    ("h", "3. 대응 방안"),
    ("p", [
        ("위기상담 인력을 더 늘리는 것만으로는 이 문제를 풀 수 없다. 사람들은 이미 그 공식 시스템의 무게를 견디지 못해 기계로 도망쳤기 때문이다. 2036년의 실패를 막기 위해 지금 당장 논의해야 할 것은, AI를 얼마나 더 사람처럼 다정하게 만들 것인가가 아니다. 기계에 갇힌 사람들을 현실의 불완전한 인간들 사이로 다시 이끌어낼 방법을 찾는 일이다.", False, False),
    ]),
    ("p", [
        ("문제는 위험을 감지하지 못하는 것이 아니라, 무엇을 위험으로 볼 것인가에 있다. 지금의 AI는 '죽고 싶다'는 명시적인 문장을 찾는다. 그러나 진짜 위기는 그런 흔적조차 남길 기력이 없는 사람들에게 더 자주 찾아온다. 단기적인 위험 방어를 넘어, 만성적인 고립을 감지하는 모델이 필요하다. 이제는 AI가 무엇을 위험으로 볼 것인지부터 다시 정의하는 일이 먼저다. ", False, False),
        ("명시적 단어가 아니라 체류 시간, 심야 대화 빈도 같은 정서적 의존 패턴을 위험 신호로 삼는 감지 기준을 1~2년 안에 설계 표준으로 세워야 한다.", True, False),
    ]),
    ("p", [
        ("이러한 의존 패턴이 감지되었을 때, ", False, False),
        ("AI의 역할은 위로를 건네는 상담사가 아니라 '연결자'가 되어야 한다.", True, False),
        (" 기계는 무작정 핫라인 번호를 띄우며 대화를 끊는 대신, 가족처럼 부담스럽지도 병원처럼 딱딱하지도 않은 지역 사회의 '제3의 공간'으로 사용자를 안내해야 한다.", False, False),
    ]),
    ("p", [
        ("거창한 대형 시설이나 기관을 새로 설립할 필요는 없다. 도서관, 청년센터, 주민센터처럼 이미 존재하는 공간이면 충분하다. AI는 이곳에서 열리는 익명 대화 부스나 가벼운 산책 모임을 자연스럽게 제안하며, 현실의 느슨한 연대로 발걸음을 돌리게 돕는 역할에 머문다.", False, False),
    ]),
    ("p", [
        ("2036년의 실패는 AI가 너무 똑똑해서 생긴 일이 아니다. 인간이 가장 외로운 순간을 더 이상 인간의 영역으로 바라보지 않았기 때문에 생긴 실패다.", False, True),
    ]),
    ("p", [
        ("앞으로 AI가 해야 할 일은 사람을 대신해 위로를 끝맺는 존재가 되는 것이 아니다. 사람이 사람에게 다시 닿을 수 있도록 첫 문장을 건네는 것이다. ", False, False),
        ("결국 인간을 회복시키는 것은 완벽한 기술이 아니라, 불완전한 관계다.", True, False),
        (" 그리고 그 누구도 위험이라 부르지 않았던 그 대화를, 이제는 우리 사회가 위험이라 이름 붙이고 책임져야 한다. ", False, False),
        ("그것이 AI 시대에 인간의 관계를 지키는 첫걸음이다.", True, False),
    ]),
    ("note", [
        ("¹ 경기연구원, 「AI를 활용한 정신건강 정책 현황 및 개선 방안」(2026); 세계일보, \"낙인 피해… 'AI 상담' 찾는 우울증 고위험군\"(2026.3.18) 재인용.", False, False),
    ]),
]

# ---------------- DOCX ----------------
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)


def set_run_font(run, font_name):
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run(TITLE)
r.font.size = Pt(18)
r.font.bold = True
set_run_font(r, HEADER_FONT)

for kind, content in SECTIONS:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.28
    if kind == "h":
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(content)
        r.font.size = Pt(12)
        r.font.bold = True
        set_run_font(r, HEADER_FONT)
    elif kind == "note":
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(0)
        for text, bold, italic in content:
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.bold = bold
            r.font.italic = italic
            set_run_font(r, BODY_FONT)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(9)
        for text, bold, italic in content:
            r = p.add_run(text)
            r.font.size = Pt(11)
            r.font.bold = bold
            r.font.italic = italic
            set_run_font(r, BODY_FONT)

doc.save(DOCX_OUT)
print("Wrote", DOCX_OUT)

# ---------------- PDF ----------------
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import simpleSplit

pdfmetrics.registerFont(TTFont("Header", HEADER_TTF))
pdfmetrics.registerFont(TTFont("Body", BODY_TTF))

W, H = A4
margin = 2.2 * cm
usable_w = W - 2 * margin
c = canvas.Canvas(str(PDF_OUT), pagesize=A4)

y = H - margin
current_font = [None, None]
c.setLineWidth(0.45)  # stroke weight used by bold (fill+stroke) text mode


def ensure_space(needed):
    global y
    if y - needed < margin:
        c.showPage()
        y = H - margin
        if current_font[0]:
            c.setFont(*current_font)


def set_font(name, size):
    c.setFont(name, size)
    current_font[0], current_font[1] = name, size


def draw_run_line(segments, size):
    """Draw one already-wrapped line made of (text, bold, italic) segments.

    Neither Chosun font has a real bold/italic face, so synthetic weight/
    slant is done via native PDF text-rendering-mode (fill+stroke for bold)
    and a sheared text matrix (for italic) -- a single Tj per run, unlike a
    double-drawString trick, which would leave two overlapping copies of
    the same text in the PDF's text layer (visible as garbled duplicates
    when copy-pasted or text-extracted).
    """
    global y
    x = margin
    for text, bold, italic in segments:
        w = pdfmetrics.stringWidth(text, "Body", size)
        if italic:
            c.saveState()
            c.translate(x, y)
            c.transform(1, 0, 0.25, 1, 0, 0)  # shear around the text's own origin
            to = c.beginText(0, 0)
        else:
            to = c.beginText(x, y)
        to.setFont("Body", size)
        # Render mode is canvas-level state that persists across text
        # objects until changed again, so it must be set explicitly on
        # every segment (not just bold ones) or a later plain segment
        # would silently inherit the previous bold's fill+stroke mode.
        if bold:
            to.setStrokeColorRGB(0, 0, 0)
            to.setTextRenderMode(2)  # fill + stroke
        else:
            to.setTextRenderMode(0)  # fill only
        to.textOut(text)
        c.drawText(to)
        if italic:
            c.restoreState()
        x += w
    y -= size * 1.42


_TOKEN_RE = re.compile(r"[A-Za-z0-9]+|.", re.DOTALL)


def wrap_runs(runs, size, max_w):
    """Word-wrap a list of (text, bold, italic) runs into lines of
    (text, bold, italic) segments. Korean text can break between any two
    syllables (normal CJK behavior), but ASCII runs like "AI" or "2036"
    are kept atomic so they never split mid-token."""
    tokens = []
    for text, bold, italic in runs:
        for tok in _TOKEN_RE.findall(text):
            tokens.append((tok, bold, italic))
    lines = []
    line = []
    line_w = 0.0
    for tok, bold, italic in tokens:
        w = pdfmetrics.stringWidth(tok, "Body", size)
        if tok == "\n":
            lines.append(line)
            line = []
            line_w = 0.0
            continue
        if line_w + w > max_w and line:
            lines.append(line)
            line = []
            line_w = 0.0
        line.append((tok, bold, italic))
        line_w += w
    if line:
        lines.append(line)
    # Merge consecutive same-style chars back into segments per line
    out_lines = []
    for line in lines:
        segs = []
        for ch, bold, italic in line:
            if segs and segs[-1][1] == bold and segs[-1][2] == italic:
                segs[-1] = (segs[-1][0] + ch, bold, italic)
            else:
                segs.append((ch, bold, italic))
        out_lines.append(segs)
    return out_lines


def draw_paragraph(runs, size=10.5, gap_before=4):
    global y
    ensure_space(gap_before + size * 1.42)
    y -= gap_before
    for line_segs in wrap_runs(runs, size, usable_w):
        ensure_space(size * 1.42)
        draw_run_line(line_segs, size)


def draw_header(text, size=13, gap_before=10):
    global y
    ensure_space(gap_before + size * 1.42)
    y -= gap_before
    set_font("Header", size)
    c.drawString(margin, y, text)
    y -= size * 1.5


# Title
set_font("Header", 18)
tw = c.stringWidth(TITLE, "Header", 18)
ensure_space(24)
c.drawString((W - tw) / 2, y, TITLE)
y -= 30

for kind, content in SECTIONS:
    if kind == "h":
        draw_header(content)
    elif kind == "note":
        draw_paragraph(content, size=8, gap_before=14)
    else:
        draw_paragraph(content)

c.showPage()
c.save()
print("Wrote", PDF_OUT)
