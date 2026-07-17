# -*- coding: utf-8 -*-
"""Regenerate 제안서_수정본.docx / .pdf from 제안서_최종본.txt (A4, <=2 pages)."""
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "제안서_최종본.txt"
DOCX_OUT = ROOT / "제안서_수정본.docx"
PDF_OUT = ROOT / "제안서_수정본.pdf"

text = SRC.read_text(encoding="utf-8")
blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
title = blocks[0]
body_blocks = blocks[1:]

FONT_KO = "Malgun Gothic"

# ---------- DOCX ----------
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

normal = doc.styles["Normal"]
normal.font.name = FONT_KO
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)

def add_para(text_, size=10.5, bold=False, space_after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.28
    if align:
        p.alignment = align
    run = p.add_run(text_)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_KO
    run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    return p

add_para(title, size=17, bold=True, space_after=16, align=WD_ALIGN_PARAGRAPH.CENTER)

for block in body_blocks:
    m = re.match(r"^(\d\. .+)$", block.splitlines()[0])
    lines = block.splitlines()
    if re.match(r"^\d\. ", lines[0]):
        add_para(lines[0], size=12.5, bold=True, space_after=6)
        rest = "\n".join(lines[1:]).strip()
        if rest:
            add_para(rest, size=10.5, space_after=10)
    else:
        add_para(block, size=10.5, space_after=10)

doc.save(DOCX_OUT)
print("Wrote", DOCX_OUT)

# ---------- PDF ----------
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import simpleSplit

import glob as _glob
candidates = [
    r"C:\Windows\Fonts\malgun.ttf",
    r"C:\Windows\Fonts\NanumGothic.ttf",
]
font_path = next((c for c in candidates if Path(c).exists()), None)
if not font_path:
    raise SystemExit("No Korean TTF font found on system (checked malgun.ttf).")

pdfmetrics.registerFont(TTFont("Korean", font_path))

W, H = A4
margin = 2.2 * cm
usable_w = W - 2 * margin
c = canvas.Canvas(str(PDF_OUT), pagesize=A4)

y = H - margin
current_font_size = [None, None]

def ensure_space(needed):
    global y
    if y - needed < margin:
        c.showPage()
        y = H - margin
        # showPage() resets ReportLab's graphics state, including the font.
        # Without this, any text drawn after a mid-paragraph page break
        # falls back to Helvetica, which has no Korean glyphs (renders as
        # solid .notdef boxes instead of the characters).
        if current_font_size[0]:
            c.setFont(*current_font_size)

def draw_wrapped(text_, size, leading, bold_gap_before=0, indent=0):
    global y
    ensure_space(bold_gap_before)
    y -= bold_gap_before
    c.setFont("Korean", size)
    current_font_size[0], current_font_size[1] = "Korean", size
    max_w = usable_w - indent
    for para_line in text_.split("\n"):
        lines = simpleSplit(para_line, "Korean", size, max_w) if para_line else [""]
        for ln in lines:
            ensure_space(leading)
            c.drawString(margin + indent, y, ln)
            y -= leading

# Title
c.setFont("Korean", 17)
tw = c.stringWidth(title, "Korean", 17)
ensure_space(24)
c.drawString((W - tw) / 2, y, title)
y -= 34

for block in body_blocks:
    lines = block.splitlines()
    if re.match(r"^\d\. ", lines[0]):
        draw_wrapped(lines[0], 13, 18, bold_gap_before=4)
        rest = "\n".join(lines[1:]).strip()
        if rest:
            draw_wrapped(rest, 10.5, 15.5, bold_gap_before=4)
    else:
        draw_wrapped(block, 10.5, 15.5, bold_gap_before=4)

c.showPage()
c.save()
print("Wrote", PDF_OUT)
